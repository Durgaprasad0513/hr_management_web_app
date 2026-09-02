import os
import sys

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def create_report():
    doc = Document()
    doc.add_heading('HR Management Web App - Daily Progress Report', 0)

    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Today, the HR Management Web App frontend was completely overhauled to match the 15-page Visily "HRMagnet" UI design template. All pages and core components were restyled to use the new layout, color schemes, and structural designs requested.')

    doc.add_heading('Key Accomplishments', level=1)

    doc.add_heading('1. Design Foundation & Layout', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Global Theme: ').bold = True
    p1.add_run('Updated Tailwind configuration with Visily theme tokens (Navy #1B2559 and Orange Accent #F97316).\n')
    p1.add_run('Navigation: ').bold = True
    p1.add_run('Removed the legacy left sidebar and implemented a sleek, horizontal Top Navigation Bar with dropdowns and quick actions.\n')
    p1.add_run('Core Components: ').bold = True
    p1.add_run('Redesigned Buttons (pill-shaped), Cards (rounded with soft shadows), Inputs, Badges, and significantly upgraded the DataTable to support checkboxes and pagination.')

    doc.add_heading('2. Complete Page Redesigns', level=2)
    doc.add_paragraph('Login Page: Split-screen design with the promotional "Employee Engagement Hub" graphics.', style='List Bullet')
    doc.add_paragraph('Dashboard: Rich widget grid displaying upcoming events, total employee statistics (donut chart), recruitment pipelines, and pending approvals.', style='List Bullet')
    doc.add_paragraph('Employee Management: Tabbed interface featuring a standard list view and a new "Directory" card grid view.', style='List Bullet')
    doc.add_paragraph('Employee Profile: Modern scroll-spy layout with a sticky right-hand navigation menu for General, Job, Time Off, and Payroll sections.', style='List Bullet')
    doc.add_paragraph('Time Tracking: Replaced the attendance module with "My Time Tracking" summaries and "Team Time Tracking" for managers to approve/reject timesheets.', style='List Bullet')
    doc.add_paragraph('Performance: Transformed into a vertical timeline view grouped by year.', style='List Bullet')
    doc.add_paragraph('Recruitment: Implemented a tabbed interface (Jobs, Candidates, Working Process) with an updated Kanban board.', style='List Bullet')
    doc.add_paragraph('Documents: Converted policies into a searchable data table with file icons, download buttons, and share links.', style='List Bullet')
    doc.add_paragraph('Company News: Created a brand new page featuring a feed of company announcements with category icons.', style='List Bullet')
    doc.add_paragraph('Task Lists: Created a new page separating To-Do, Onboarding, and Offboarding tasks with progress bars.', style='List Bullet')

    doc.add_heading('3. Code Quality & Testing', level=2)
    doc.add_paragraph('Successfully ran the TypeScript compiler across the entire frontend. Resolved minor syntax issues (e.g., JSX template literal escaping in the Task List Page), resulting in a clean, zero-error build.')

    save_path = r'c:\Users\VDurgaprasad\OneDrive\Desktop\HR_Web_App_Daily_Report.docx'
    doc.save(save_path)
    print(f"Report saved to {save_path}")

if __name__ == '__main__':
    create_report()
